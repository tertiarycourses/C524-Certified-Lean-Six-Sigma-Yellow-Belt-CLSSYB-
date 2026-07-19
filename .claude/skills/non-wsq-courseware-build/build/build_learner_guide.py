#!/usr/bin/env python3
"""Generate the non-WSQ Learner Guide as BOTH a Markdown
mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, wrap-up and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
# Dynamic domain import — supports ANY number of data_domainN.py files.
import importlib, glob as _g0
def _load_domains():
    acts=[]
    for f in sorted(_g0.glob(os.path.join(HERE,"data_domain[0-9]*.py")),
                    key=lambda q:int("".join(c for c in os.path.basename(q) if c.isdigit()) or 0)):
        n="".join(c for c in os.path.basename(f) if c.isdigit())
        acts+=getattr(importlib.import_module(os.path.basename(f)[:-3]),f"DOMAIN{n}",[])
    return acts
ACT=_load_domains()
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# Pull the DETAILED step-by-step from each lab's README so the Learner Guide is
# highly detailed AND stays 100% aligned with the labs/ folder. Falls back to the
# terse course_data steps if a README is missing/unparseable.
import glob as _glob, re as _re
def _readme_steps(num):
    cands=_glob.glob(os.path.join(REPO,"labs",f"lab{num:02d}-*","README.md"))
    if not cands: return None
    txt=open(cands[0],encoding="utf-8").read()
    # Terminate the Steps body only at a 2-hash section heading ("## Verification"),
    # NOT at 3-hash "### N." step headers used by some labs.
    m=_re.search(r'\n##\s*Steps\s*\n(.*?)(?=\n##\s|\Z)', txt, _re.S)
    if not m: return None
    body=m.group(1).strip()
    # Handle both step formats: "N. **Title.** prose" list items AND "### N. Title" headers.
    items=_re.split(r'\n(?=(?:#{2,4}\s*)?\d+\.\s)', body)
    out=[]
    for it in items:
        it=it.strip()
        if not _re.match(r'(?:#{2,4}\s*)?\d+\.\s', it): continue
        cm=_re.search(r'```[a-zA-Z0-9]*\n(.*?)```', it, _re.S)
        cmd=cm.group(1).strip() if cm else ""
        # ignore non-command fenced snippets (e.g. ```text chat prompts) — keep only real commands
        if cmd and _re.match(r'^\s*#', cmd) and "\n" not in cmd:
            cmd=""
        text=it[:cm.start()] if cm else it
        # drop markdown table rows (| ... |) — tables don't flatten into prose
        text="\n".join(l for l in text.splitlines() if not l.strip().startswith("|"))
        text=_re.sub(r'^(?:#{2,4}\s*)?\d+\.\s*','',text)      # drop leading '### N.' / 'N.'
        text=_re.sub(r'\*\*(.*?)\*\*', r'\1', text)           # drop bold markers
        text=_re.sub(r'\*(.*?)\*', r'\1', text)               # drop italics
        text=_re.sub(r'`([^`]+)`', r'\1', text)               # drop inline code ticks
        text=_re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)  # link -> text (url)
        text=_re.sub(r'\s*\n\s*',' ', text).strip()
        if text:
            out.append((text,cmd))
    return out or None

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(getattr(C,"LG_INTRO",
  f"This Learner Guide accompanies the course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"It provides step-by-step instructions for all {len(ACT)} hands-on labs, organised into "
  f"{len(C.TOPICS)} topics that follow the course slides and Lesson Plan."))
p(getattr(C,"LG_INTRO2",
  "Use this guide alongside the course slides and the lab files in the labs/ folder of the course "
  "repository. Work through the labs in order — each one builds on the skills established by the "
  "labs before it."))

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

_SETUP=getattr(C,"LG_SETUP",None)
if _SETUP:
    h1("Before You Start — Preparation")
    if _SETUP.get("needs"):
        h3("What you need"); bullets(_SETUP["needs"])
    if _SETUP.get("verify_text") or _SETUP.get("verify_code"):
        h3("Verify your setup")
        if _SETUP.get("verify_text"): p(_SETUP["verify_text"])
        if _SETUP.get("verify_code"): code(_SETUP["verify_code"])
    if _SETUP.get("conventions"):
        h3("Conventions used in every lab"); bullets(_SETUP["conventions"])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}" + (f"  ({t['weighting']})" if t.get("weighting") else ""))
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Learning outcome: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Step-by-step")
        # Prefer the detailed README steps. The LG never shows YouTube reference
        # links — videos live in the labs/ READMEs only.
        st=_readme_steps(a["num"]) or [(instr,cmd) for (instr,cmd) in a["steps"]]
        st=[(si,sc) for (si,sc) in st if "youtube" not in sc.lower() and "youtube" not in si.lower()]
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"Full commands and screenshots are in labs/lab-{a['num']:02d}-*.md. "
             + getattr(C,"LAB_NOTE","Use only accounts and data you are authorised to use."))
        rule()

_WRAP=getattr(C,"LG_WRAPUP",None)
if _WRAP:
    h1(_WRAP.get("title","Wrap-Up"))
    if _WRAP.get("intro"): p(_WRAP["intro"])
    for sec in _WRAP.get("sections",[]):
        h3(sec.get("title","")); 
        if sec.get("text"): p(sec["text"])
        if sec.get("bullets"): bullets(sec["bullets"])
    rule()

h1("Next Steps")
bullets(getattr(C,"LG_NEXT_STEPS",[
 "First pass: complete every lab yourself, following the steps in each lab file.",
 "Second pass: redo the labs from memory until the workflow is automatic.",
 "Apply the techniques to a real process or project in your own organisation.",
 "Review each lab's detailed steps in this guide and re-run the labs on your own.",
]))

_GLOSSARY=getattr(C,"LG_GLOSSARY",[])
if _GLOSSARY:
    h1("Glossary")
    B.append(("dl",_GLOSSARY))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,getattr(C,"VERSION_HISTORY",
 [(C.VERSION.lstrip("v"),C.VERSION_DATE,"Initial release.",C.TRAINER)]))
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            # Literal per-lab step numbers (Word's List Number style numbers
            # continuously across the whole document — 1..N must restart per lab).
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.space_after=Pt(4)
            r=para.add_run(f"{i}.  "); r.bold=True
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
